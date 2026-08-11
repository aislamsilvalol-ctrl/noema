"""PDF, DOCX and web page parsers.

Each third-party library is imported inside its function. Parsers are the largest
untrusted-input surface in the system, and a heavy optional dependency should not
decide whether the API process can start.
"""

from __future__ import annotations

import re
from typing import Any

from noema.core.errors import NoemaError
from noema.ingestion.ir import Block, BlockKind, ParsedDocument

#: A page yielding less than this is treated as scanned and sent to OCR.
OCR_CHAR_THRESHOLD = 50

#: PDFs have no heading markup, so headings are inferred from relative font size.
HEADING_SIZE_RATIO = 1.15


class ParserUnavailable(NoemaError):
    slug = "parser-unavailable"
    title = "Parser unavailable"
    status_code = 501


class ParseFailed(NoemaError):
    slug = "parse-failed"
    title = "Could not read this document"
    status_code = 422


def parse_pdf(data: bytes, *, ocr: bool = True) -> ParsedDocument:
    """Extract text with page anchors, inferring headings from font size."""
    try:
        import pymupdf
    except ImportError as exc:  # pragma: no cover - exercised only without the extra
        raise ParserUnavailable("PDF support requires pymupdf") from exc

    try:
        document: Any = pymupdf.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ParseFailed("This file could not be opened as a PDF.") from exc

    blocks: list[Block] = []
    body_size = _median_font_size(document)

    with document:
        for index in range(document.page_count):
            number = index + 1
            page: Any = document[index]
            page_blocks = _blocks_from_page(page, number, body_size)

            if ocr and _visible_chars(page_blocks) < OCR_CHAR_THRESHOLD:
                page_blocks = _ocr_page(page, number) or page_blocks

            blocks.extend(page_blocks)

        metadata: dict[str, Any] = {
            key: value
            for key, value in (document.metadata or {}).items()
            if value and key in {"title", "author", "subject", "keywords"}
        }
        page_count = document.page_count

    return ParsedDocument(blocks=blocks, metadata=metadata, page_count=page_count)


def parse_docx(data: bytes) -> ParsedDocument:
    """DOCX keeps real heading levels, so the structure survives intact."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ParserUnavailable("DOCX support requires python-docx") from exc

    import io

    try:
        document: Any = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ParseFailed("This file could not be opened as a Word document.") from exc

    blocks: list[Block] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else None
        style = (style_name or "").lower()
        if style.startswith("heading"):
            level = _heading_level(style)
            blocks.append(Block(kind=BlockKind.HEADING, text=text, level=level))
        elif style.startswith("list"):
            blocks.append(Block(kind=BlockKind.LIST_ITEM, text=text))
        elif style.startswith("quote"):
            blocks.append(Block(kind=BlockKind.QUOTE, text=text))
        else:
            blocks.append(Block(kind=BlockKind.PARAGRAPH, text=text))

    for table in document.tables:
        rows = [", ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if any(rows):
            blocks.append(Block(kind=BlockKind.TABLE, text="\n".join(rows)))

    return ParsedDocument(blocks=blocks)


def parse_html(html: str, url: str | None = None) -> ParsedDocument:
    """Strip boilerplate from a web page and keep the article."""
    try:
        import trafilatura
    except ImportError as exc:  # pragma: no cover
        raise ParserUnavailable("URL support requires trafilatura") from exc

    extracted = trafilatura.extract(
        html, include_comments=False, include_tables=True, output_format="markdown"
    )
    if not extracted:
        raise ParseFailed("No readable article was found at this address.")

    from noema.ingestion.parsers.text import parse_markdown

    document = parse_markdown(extracted)
    metadata = dict(document.metadata)
    if url:
        metadata["url"] = url

    extracted_meta = trafilatura.extract_metadata(html)
    if extracted_meta is not None:
        for key in ("title", "author", "date"):
            value = getattr(extracted_meta, key, None)
            if value:
                metadata.setdefault(key, value)

    return ParsedDocument(blocks=document.blocks, metadata=metadata)


def parse_transcript(text: str) -> ParsedDocument:
    """WebVTT or SRT. Timestamps are kept as citation anchors."""
    cue_time = re.compile(r"(\d{2}):(\d{2}):(\d{2})[.,](\d{3})\s*-->")
    blocks: list[Block] = []
    current: list[str] = []
    start: float | None = None

    def flush() -> None:
        joined = " ".join(current).strip()
        if joined:
            blocks.append(Block(kind=BlockKind.PARAGRAPH, text=joined, timestamp=start))
        current.clear()

    for line in text.replace("\r\n", "\n").split("\n"):
        stripped = line.strip()
        match = cue_time.match(stripped)
        if match:
            flush()
            hours, minutes, seconds, millis = (int(g) for g in match.groups())
            start = hours * 3600 + minutes * 60 + seconds + millis / 1000
            continue
        if not stripped or stripped.isdigit() or stripped.upper() == "WEBVTT":
            continue
        current.append(stripped)

    flush()
    return ParsedDocument(blocks=blocks)


def _blocks_from_page(page: Any, number: int, body_size: float) -> list[Block]:
    blocks: list[Block] = []
    data = page.get_text("dict")

    for block in data.get("blocks", []):
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            text = "".join(span.get("text", "") for span in spans).strip()
            if not text:
                continue

            size = max((span.get("size", 0.0) for span in spans), default=0.0)
            is_heading = body_size > 0 and size >= body_size * HEADING_SIZE_RATIO

            if is_heading:
                # Two sizes above body text reads as a section, one as a subsection.
                level = 1 if size >= body_size * 1.4 else 2
                blocks.append(
                    Block(kind=BlockKind.HEADING, text=text, level=level, page=number)
                )
            else:
                blocks.append(Block(kind=BlockKind.PARAGRAPH, text=text, page=number))

    return _join_wrapped_lines(blocks)


def _join_wrapped_lines(blocks: list[Block]) -> list[Block]:
    """PDF extraction yields one block per visual line; rejoin them into paragraphs.

    Without this, every chunk boundary lands mid-sentence and citations quote half a
    clause.
    """
    joined: list[Block] = []
    for block in blocks:
        previous = joined[-1] if joined else None
        continues = (
            previous is not None
            and previous.kind is BlockKind.PARAGRAPH
            and block.kind is BlockKind.PARAGRAPH
            and previous.page == block.page
            and not previous.text.endswith((".", "!", "?", ":", ";"))
        )
        if continues and previous is not None:
            separator = "" if previous.text.endswith("-") else " "
            joined[-1] = Block(
                kind=BlockKind.PARAGRAPH,
                text=previous.text.rstrip("-") + separator + block.text,
                page=previous.page,
            )
        else:
            joined.append(block)
    return joined


def _median_font_size(document: Any) -> float:
    sizes: list[float] = []
    for page in document.pages(0, min(5, document.page_count)):
        for block in page.get_text("dict").get("blocks", []):
            for line in block.get("lines", []):
                sizes.extend(span.get("size", 0.0) for span in line.get("spans", []))
    if not sizes:
        return 0.0
    sizes.sort()
    return sizes[len(sizes) // 2]


def _visible_chars(blocks: list[Block]) -> int:
    return sum(len(block.text.strip()) for block in blocks)


def _ocr_page(page: Any, number: int) -> list[Block]:
    """OCR fallback for scanned pages. Returns nothing if Tesseract is unavailable."""
    try:
        textpage = page.get_textpage_ocr(flags=0, full=True)
    except Exception:
        return []

    text = page.get_text(textpage=textpage).strip()
    if not text:
        return []

    return [
        Block(kind=BlockKind.PARAGRAPH, text=paragraph.strip(), page=number)
        for paragraph in re.split(r"\n\s*\n", text)
        if paragraph.strip()
    ]


def _heading_level(style_name: str) -> int:
    digits = re.search(r"(\d+)", style_name)
    return min(max(int(digits.group(1)), 1), 6) if digits else 1
